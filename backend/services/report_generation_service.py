"""
Report Generation Service - BlindSpot.AI v2
Generates professional PDF and DOCX evaluation reports.
"""
import os, json, logging
from datetime import datetime
from typing import Dict, Any, List

logger = logging.getLogger(__name__)
DATA_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "data")


def generate_report(evaluation_id: str, evaluation_data: Dict[str, Any]) -> str:
    report = _build_report_dict(evaluation_id, evaluation_data)
    out_dir = os.path.join(DATA_DIR, "reports", evaluation_id)
    os.makedirs(out_dir, exist_ok=True)

    with open(os.path.join(out_dir, "report.json"), "w", encoding="utf-8") as f:
        json.dump(report, f, indent=2, default=str)

    pdf_path  = os.path.join(out_dir, "report.pdf")
    docx_path = os.path.join(out_dir, "report.docx")

    try:
        _build_pdf(report, pdf_path)
    except Exception as e:
        logger.error(f"[ReportGen] PDF failed: {e}")

    try:
        _build_docx(report, docx_path)
    except Exception as e:
        logger.error(f"[ReportGen] DOCX failed: {e}")

    logger.info(f"[ReportGen] Reports saved to {out_dir}")
    return f"http://localhost:8000/media/reports/{evaluation_id}/report.pdf"


# ─── Report dict builder ──────────────────────────────────────────────────────

def _build_report_dict(evaluation_id: str, d: Dict[str, Any]) -> Dict[str, Any]:
    stress  = d.get("stress_results") or []
    datasets= d.get("fetched_datasets") or []
    edges   = d.get("edge_case_analysis") or []
    metrics = d.get("original_metrics") or {}
    weakness= d.get("weakness_report") or {}
    score   = float(d.get("robustness_score") or 0)
    return {
        "report_id":    f"BSR-{evaluation_id[:8].upper()}",
        "generated_at": datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC"),
        "platform":     "BlindSpot.AI v2.0",
        "model": {
            "name":         d.get("name","Unnamed"),
            "architecture": d.get("architecture","Unknown"),
            "framework":    d.get("framework","Unknown"),
            "dataset_type": d.get("dataset_type","Unknown"),
            "model_file":   d.get("model_filename","N/A"),
            "metrics": {k: v for k, v in metrics.items() if v is not None},
        },
        "scope": {
            "task_type":  d.get("detected_task_type","Unknown"),
            "domain":     d.get("domain","Unknown"),
            "summary":    d.get("scope_summary",""),
        },
        "edge_cases": {
            "total":    len(edges),
            "critical": sum(1 for e in edges if e.get("severity")=="critical"),
            "high":     sum(1 for e in edges if e.get("severity")=="high"),
            "medium":   sum(1 for e in edges if e.get("severity")=="medium"),
            "items":    edges,
            "weaknesses":   weakness.get("weaknesses",[]),
            "risk_factors": weakness.get("risk_factors",[]),
        },
        "datasets": {
            "total":   len(datasets),
            "samples": sum((ds.get("sample_count") or ds.get("samples",0)) for ds in datasets),
            "items":   datasets,
        },
        "stress_tests": {
            "total":   len(stress),
            "passed":  sum(1 for r in stress if r.get("passed")),
            "failed":  sum(1 for r in stress if not r.get("passed")),
            "avg_deg": round(sum(float(r.get("degradation_pct") or 0) for r in stress)/len(stress),1) if stress else 0,
            "worst":   min(stress, key=lambda r: float(r.get("stressed_score") or 1), default={}).get("stressor_label","N/A"),
            "items":   stress,
        },
        "assessment": {
            "robustness_score": score,
            "risk_level":       d.get("risk_level","unknown"),
            "deployment_ready": d.get("deployment_ready",False),
            "recommendation":   _recommendation(score, d.get("deployment_ready",False)),
            "action_items":     _action_items(stress, edges),
        },
    }


def _recommendation(score: float, ready: bool) -> str:
    if ready and score >= 80:
        return (f"Model demonstrates strong robustness ({score:.1f}%). "
                "Approved for production deployment with standard monitoring.")
    elif score >= 60:
        return (f"Model shows moderate robustness ({score:.1f}%). "
                "Conditional deployment approved. Augment training data for failed stressors.")
    else:
        return (f"Robustness score {score:.1f}% is below production threshold. "
                "Deployment NOT recommended. Retrain with stress-augmented datasets.")


def _action_items(stress: List[Dict], edges: List[Dict]) -> List[str]:
    items = []
    for r in [r for r in (stress or []) if r and not r.get("passed")][:3]:
        label = r.get("stressor_label") or r.get("stressor_key","unknown")
        score = float(r.get("stressed_score") or 0) * 100
        items.append(f"Augment training with {label} data (stressed accuracy: {score:.1f}%)")
    for ec in [e for e in (edges or []) if e and e.get("severity") == "critical"][:2]:
        items.append(f"Address critical edge case: {ec.get('name','?')}")
    if not items:
        items.append("Monitor model confidence thresholds in production.")
    return items


# ─── PDF Builder ─────────────────────────────────────────────────────────────

def _build_pdf(report: Dict[str, Any], path: str):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable, KeepTogether)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT

    doc = SimpleDocTemplate(path, pagesize=A4,
                            leftMargin=2*cm, rightMargin=2*cm,
                            topMargin=2*cm, bottomMargin=2*cm)

    # ── Colour palette ────────────────────────────────────────────────────
    DARK   = colors.HexColor("#0d1117")
    BLUE   = colors.HexColor("#3b82f6")
    EMERALD= colors.HexColor("#10b981")
    AMBER  = colors.HexColor("#f59e0b")
    RED    = colors.HexColor("#ef4444")
    SLATE  = colors.HexColor("#64748b")
    WHITE  = colors.white
    LIGHT  = colors.HexColor("#f1f5f9")

    styles = getSampleStyleSheet()

    def sty(name, **kw):
        return ParagraphStyle(name, parent=styles["Normal"], **kw)

    H1   = sty("H1",   fontSize=22, textColor=DARK,   spaceAfter=4,  spaceBefore=12, fontName="Helvetica-Bold")
    H2   = sty("H2",   fontSize=14, textColor=BLUE,   spaceAfter=4,  spaceBefore=10, fontName="Helvetica-Bold")
    H3   = sty("H3",   fontSize=11, textColor=DARK,   spaceAfter=3,  spaceBefore=6,  fontName="Helvetica-Bold")
    BODY = sty("BODY", fontSize=9,  textColor=SLATE,  spaceAfter=3,  leading=14)
    MONO = sty("MONO", fontSize=8,  textColor=DARK,   fontName="Courier", spaceAfter=2)
    CENT = sty("CENT", fontSize=9,  textColor=SLATE,  alignment=TA_CENTER)
    WARN = sty("WARN", fontSize=9,  textColor=RED,    fontName="Helvetica-Bold")
    OK   = sty("OK",   fontSize=9,  textColor=EMERALD,fontName="Helvetica-Bold")

    def score_color(s):
        if s >= 80: return EMERALD
        if s >= 60: return AMBER
        return RED

    story = []

    # ── Cover ─────────────────────────────────────────────────────────────
    story.append(Spacer(1, 1*cm))
    story.append(Paragraph("BlindSpot.AI", sty("brand", fontSize=11, textColor=BLUE, fontName="Helvetica-Bold")))
    story.append(Paragraph("Model Robustness Evaluation Report", H1))
    story.append(HRFlowable(width="100%", thickness=2, color=BLUE, spaceAfter=6))

    m = report["model"]
    a = report["assessment"]
    score = float(a["robustness_score"])

    # Header info table
    hdr_data = [
        ["Report ID",       report["report_id"],    "Generated",    report["generated_at"]],
        ["Model",           m["name"],              "Framework",    m["framework"]],
        ["Architecture",    m["architecture"],      "Dataset Type", m["dataset_type"]],
        ["Task Type",       report["scope"]["task_type"], "Domain", report["scope"]["domain"]],
    ]
    hdr_tbl = Table(hdr_data, colWidths=[3.5*cm, 6*cm, 3.5*cm, 4*cm])
    hdr_tbl.setStyle(TableStyle([
        ("FONTNAME",  (0,0),(-1,-1), "Helvetica"),
        ("FONTSIZE",  (0,0),(-1,-1), 8),
        ("FONTNAME",  (0,0),(0,-1),  "Helvetica-Bold"),
        ("FONTNAME",  (2,0),(2,-1),  "Helvetica-Bold"),
        ("TEXTCOLOR", (0,0),(0,-1),  SLATE),
        ("TEXTCOLOR", (2,0),(2,-1),  SLATE),
        ("TEXTCOLOR", (1,0),(1,-1),  DARK),
        ("TEXTCOLOR", (3,0),(3,-1),  DARK),
        ("ROWBACKGROUNDS", (0,0),(-1,-1), [LIGHT, WHITE]),
        ("GRID",      (0,0),(-1,-1), 0.3, colors.HexColor("#e2e8f0")),
        ("PADDING",   (0,0),(-1,-1), 5),
    ]))
    story.append(hdr_tbl)
    story.append(Spacer(1, 0.5*cm))

    # ── Robustness Score Banner ────────────────────────────────────────────
    sc = score_color(score)
    # Convert reportlab color to 6-char hex string for XML markup
    sc_hex = sc.hexval().replace("0x", "").replace("0X", "").zfill(6)
    risk = a["risk_level"].upper()
    ready_txt = "YES - DEPLOYMENT APPROVED" if a["deployment_ready"] else "NO - NOT RECOMMENDED"
    ready_col = EMERALD if a["deployment_ready"] else RED

    banner_data = [
        [Paragraph("<b>Robustness Score</b>", CENT),
         Paragraph("<b>Risk Level</b>", CENT),
         Paragraph("<b>Deployment Ready</b>", CENT)],
        [Paragraph(f"<font color='#{sc_hex}' size='20'><b>{score:.1f}%</b></font>", CENT),
         Paragraph(f"<b>{risk}</b>", CENT),
         Paragraph(f"<b>{ready_txt}</b>", CENT)],
    ]
    banner = Table(banner_data, colWidths=[5.5*cm, 5.5*cm, 6*cm])
    banner.setStyle(TableStyle([
        ("BACKGROUND",  (0,0),(-1,0), DARK),
        ("TEXTCOLOR",   (0,0),(-1,0), WHITE),
        ("FONTNAME",    (0,0),(-1,-1),"Helvetica-Bold"),
        ("FONTSIZE",    (0,0),(-1,0), 9),
        ("ALIGN",       (0,0),(-1,-1),"CENTER"),
        ("VALIGN",      (0,0),(-1,-1),"MIDDLE"),
        ("ROWHEIGHT",   (0,0),(-1,-1), 28),
        ("GRID",        (0,0),(-1,-1), 0.5, colors.HexColor("#334155")),
        ("PADDING",     (0,0),(-1,-1), 8),
    ]))
    story.append(banner)
    story.append(Spacer(1, 0.6*cm))

    # ── Section 1: Original Metrics ───────────────────────────────────────
    story.append(Paragraph("1. Original Model Metrics", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    met = m.get("metrics", {})
    if met:
        met_rows = [["Metric", "Value", "Status"]]
        thresholds = {"accuracy":0.90,"precision":0.85,"recall":0.85,"f1":0.87,"map":0.50,"roc_auc":0.85}
        for k, v in met.items():
            if v is None: continue
            thresh = thresholds.get(k, 0.80)
            status = "PASS" if float(v) >= thresh else "WARN"
            met_rows.append([k.upper().replace("_"," "), f"{float(v):.3f}", status])
        mt = Table(met_rows, colWidths=[5*cm, 4*cm, 4*cm])
        mt.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0), DARK),
            ("TEXTCOLOR",   (0,0),(-1,0), WHITE),
            ("FONTNAME",    (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 9),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT,WHITE]),
            ("GRID",        (0,0),(-1,-1), 0.3, colors.HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 6),
        ]))
        story.append(mt)
    story.append(Spacer(1, 0.4*cm))

    # ── Section 2: Scope Analysis ─────────────────────────────────────────
    story.append(Paragraph("2. Scope Analysis", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    story.append(Paragraph(report["scope"]["summary"], BODY))
    story.append(Spacer(1, 0.4*cm))

    # ── Section 3: Edge Case Analysis ────────────────────────────────────
    ec_data = report["edge_cases"]
    story.append(Paragraph("3. Edge Case Analysis", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    story.append(Paragraph(
        f"Identified <b>{ec_data['total']}</b> edge cases: "
        f"<font color='#ef4444'>{ec_data['critical']} critical</font>, "
        f"<font color='#f59e0b'>{ec_data['high']} high</font>, "
        f"{ec_data['medium']} medium.", BODY))
    story.append(Spacer(1, 0.2*cm))

    if ec_data["items"]:
        ec_rows = [["Edge Case", "Severity", "Stressor", "Description"]]
        for ec in ec_data["items"]:
            sev = ec.get("severity","").upper()
            ec_rows.append([
                ec.get("name",""),
                sev,
                ec.get("stressor",""),
                ec.get("description","")[:80] + ("..." if len(ec.get("description",""))>80 else ""),
            ])
        ec_tbl = Table(ec_rows, colWidths=[4*cm, 2*cm, 3*cm, 8*cm])
        sev_colors = []
        for i, ec in enumerate(ec_data["items"], 1):
            c = RED if ec.get("severity")=="critical" else AMBER if ec.get("severity")=="high" else EMERALD
            sev_colors.append(("TEXTCOLOR",(1,i),(1,i),c))
        ec_tbl.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0), DARK),
            ("TEXTCOLOR",   (0,0),(-1,0), WHITE),
            ("FONTNAME",    (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT,WHITE]),
            ("GRID",        (0,0),(-1,-1), 0.3, colors.HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 5),
            ("FONTNAME",    (1,1),(1,-1), "Helvetica-Bold"),
        ] + sev_colors))
        story.append(ec_tbl)
    story.append(Spacer(1, 0.4*cm))

    # ── Section 4: Stress Test Results ───────────────────────────────────
    st_data = report["stress_tests"]
    story.append(Paragraph("4. Stress Test Results", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    story.append(Paragraph(
        f"<b>{st_data['passed']}</b> passed / <b>{st_data['failed']}</b> failed out of "
        f"<b>{st_data['total']}</b> tests. Average degradation: <b>{st_data['avg_deg']}%</b>. "
        f"Worst stressor: <b>{st_data['worst']}</b>.", BODY))
    story.append(Spacer(1, 0.2*cm))

    if st_data["items"]:
        st_rows = [["Stressor", "Original", "Stressed", "Degradation", "Stability", "Status"]]
        for r in st_data["items"]:
            if not r:
                continue
            st_rows.append([
                r.get("stressor_label") or r.get("stressor_key",""),
                f"{float(r.get('original_score') or 0)*100:.1f}%",
                f"{float(r.get('stressed_score') or 0)*100:.1f}%",
                f"{float(r.get('degradation_pct') or 0):.1f}%",
                f"{float(r.get('confidence_stability') or 0):.3f}",
                "PASS" if r.get("passed") else "FAIL",
            ])
        pass_fail_colors = []
        for i, r in enumerate(st_data["items"], 1):
            c = EMERALD if r.get("passed") else RED
            pass_fail_colors.append(("TEXTCOLOR",(5,i),(5,i),c))
            pass_fail_colors.append(("FONTNAME",(5,i),(5,i),"Helvetica-Bold"))
        st_tbl = Table(st_rows, colWidths=[4.5*cm, 2.2*cm, 2.2*cm, 2.5*cm, 2.5*cm, 2*cm])
        st_tbl.setStyle(TableStyle([
            ("BACKGROUND",  (0,0),(-1,0), DARK),
            ("TEXTCOLOR",   (0,0),(-1,0), WHITE),
            ("FONTNAME",    (0,0),(-1,0), "Helvetica-Bold"),
            ("FONTSIZE",    (0,0),(-1,-1), 8),
            ("ROWBACKGROUNDS",(0,1),(-1,-1),[LIGHT,WHITE]),
            ("GRID",        (0,0),(-1,-1), 0.3, colors.HexColor("#e2e8f0")),
            ("PADDING",     (0,0),(-1,-1), 5),
            ("ALIGN",       (1,0),(-1,-1), "CENTER"),
        ] + pass_fail_colors))
        story.append(st_tbl)
    story.append(Spacer(1, 0.4*cm))

    # ── Section 5: Dataset Summary ────────────────────────────────────────
    ds_data = report["datasets"]
    story.append(Paragraph("5. Dataset Summary", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    story.append(Paragraph(
        f"<b>{ds_data['total']}</b> datasets used, "
        f"<b>{ds_data['samples']:,}</b> total samples.", BODY))
    story.append(Spacer(1, 0.4*cm))

    # ── Section 6: Final Assessment ───────────────────────────────────────
    story.append(Paragraph("6. Final Assessment & Recommendation", H2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=BLUE, spaceAfter=4))
    story.append(Paragraph(a["recommendation"], BODY))
    story.append(Spacer(1, 0.3*cm))

    if a["action_items"]:
        story.append(Paragraph("<b>Action Items:</b>", H3))
        for item in a["action_items"]:
            story.append(Paragraph(f"  - {item}", BODY))
    story.append(Spacer(1, 0.5*cm))

    # ── Footer ────────────────────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=1, color=BLUE, spaceAfter=4))
    story.append(Paragraph(
        f"BlindSpot.AI v2.0  |  Report {report['report_id']}  |  {report['generated_at']}  |  CONFIDENTIAL",
        sty("footer", fontSize=7, textColor=SLATE, alignment=TA_CENTER)))

    doc.build(story)
    logger.info(f"[PDF] Generated: {path} ({os.path.getsize(path)//1024} KB)")


# ─── DOCX Builder ────────────────────────────────────────────────────────────

def _build_docx(report: Dict[str, Any], path: str):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    def rgb(hex_str):
        h = hex_str.lstrip("#")
        return RGBColor(int(h[0:2],16), int(h[2:4],16), int(h[4:6],16))

    def add_heading(text, level=1, color="#0d1117"):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = rgb(color)
        return p

    def add_para(text, bold=False, color=None, size=10):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        if color:
            run.font.color.rgb = rgb(color)
        return p

    def add_table(headers, rows, col_widths=None):
        tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
        tbl.style = "Table Grid"
        # Header row
        hdr_row = tbl.rows[0]
        for i, h in enumerate(headers):
            cell = hdr_row.cells[i]
            cell.text = h
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = rgb("#ffffff")
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "0d1117")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:val"), "clear")
            tcPr.append(shd)
        # Data rows
        for ri, row_data in enumerate(rows):
            row = tbl.rows[ri+1]
            for ci, val in enumerate(row_data):
                cell = row.cells[ci]
                safe_val = str(val) if val is not None else ""
                cell.text = safe_val
                run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(safe_val)
                run.font.size = Pt(8)
                if ri % 2 == 0:
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "f1f5f9")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)
        return tbl

    m = report["model"]
    a = report["assessment"]
    score = float(a["robustness_score"])

    # ── Cover ─────────────────────────────────────────────────────────────
    title = doc.add_heading("BlindSpot.AI", 0)
    title.runs[0].font.color.rgb = rgb("#3b82f6")
    title.runs[0].font.size = Pt(11)

    h = doc.add_heading("Model Robustness Evaluation Report", 0)
    h.runs[0].font.size = Pt(22)
    h.runs[0].font.color.rgb = rgb("#0d1117")

    doc.add_paragraph()
    info_tbl = doc.add_table(rows=4, cols=4)
    info_tbl.style = "Table Grid"
    info_data = [
        ["Report ID",    str(report["report_id"] or ""),       "Generated",    str(report["generated_at"] or "")],
        ["Model",        str(m["name"] or ""),                 "Framework",    str(m["framework"] or "Unknown")],
        ["Architecture", str(m["architecture"] or "Unknown"),  "Dataset Type", str(m["dataset_type"] or "Unknown")],
        ["Task Type",    str(report["scope"]["task_type"] or "Unknown"), "Domain", str(report["scope"]["domain"] or "Unknown")],
    ]
    for ri, row_data in enumerate(info_data):
        row = info_tbl.rows[ri]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(val)
            run.font.size = Pt(9)
            if ci % 2 == 0:
                run.bold = True
                run.font.color.rgb = rgb("#64748b")

    doc.add_paragraph()

    # ── Score Banner ──────────────────────────────────────────────────────
    score_color = "#10b981" if score >= 80 else "#f59e0b" if score >= 60 else "#ef4444"
    ready_txt = "YES - APPROVED" if a["deployment_ready"] else "NO - NOT RECOMMENDED"
    banner = doc.add_paragraph()
    banner.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = banner.add_run(f"Robustness Score: {score:.1f}%  |  Risk: {a['risk_level'].upper()}  |  Deployment: {ready_txt}")
    r1.bold = True
    r1.font.size = Pt(13)
    r1.font.color.rgb = rgb(score_color)
    doc.add_paragraph()

    # ── Section 1: Metrics ────────────────────────────────────────────────
    add_heading("1. Original Model Metrics", 1, "#3b82f6")
    met = m.get("metrics", {})
    if met:
        thresholds = {"accuracy":0.90,"precision":0.85,"recall":0.85,"f1":0.87,"map":0.50,"roc_auc":0.85}
        rows = []
        for k, v in met.items():
            if v is None: continue
            try:
                thresh = thresholds.get(k, 0.80)
                rows.append([k.upper().replace("_"," "), f"{float(v):.3f}", "PASS" if float(v)>=thresh else "WARN"])
            except (TypeError, ValueError):
                continue
        if rows:
            add_table(["Metric","Value","Status"], rows)
    doc.add_paragraph()

    # ── Section 2: Scope ──────────────────────────────────────────────────
    add_heading("2. Scope Analysis", 1, "#3b82f6")
    add_para(report["scope"]["summary"])
    doc.add_paragraph()

    # ── Section 3: Edge Cases ─────────────────────────────────────────────
    ec = report["edge_cases"]
    add_heading("3. Edge Case Analysis", 1, "#3b82f6")
    add_para(f"Total: {ec['total']} | Critical: {ec['critical']} | High: {ec['high']} | Medium: {ec['medium']}", bold=True)
    items = ec.get("items") or []
    if items:
        rows = [[e.get("name",""), e.get("severity","").upper(), e.get("stressor",""),
                 (e.get("description","") or "")[:80]] for e in items if e]
        if rows:
            add_table(["Edge Case","Severity","Stressor","Description"], rows)
    doc.add_paragraph()

    # ── Section 4: Stress Tests ───────────────────────────────────────────
    st = report["stress_tests"]
    add_heading("4. Stress Test Results", 1, "#3b82f6")
    add_para(f"Passed: {st['passed']} | Failed: {st['failed']} | Avg Degradation: {st['avg_deg']}% | Worst: {st['worst']}", bold=True)
    st_items = st.get("items") or []
    if st_items:
        rows = []
        for r in st_items:
            if not r: continue
            try:
                rows.append([
                    r.get("stressor_label") or r.get("stressor_key",""),
                    f"{float(r.get('original_score') or 0)*100:.1f}%",
                    f"{float(r.get('stressed_score') or 0)*100:.1f}%",
                    f"{float(r.get('degradation_pct') or 0):.1f}%",
                    "PASS" if r.get("passed") else "FAIL"
                ])
            except (TypeError, ValueError):
                continue
        if rows:
            add_table(["Stressor","Original","Stressed","Degradation","Status"], rows)
    doc.add_paragraph()

    # ── Section 5: Datasets ───────────────────────────────────────────────
    ds = report["datasets"]
    add_heading("5. Dataset Summary", 1, "#3b82f6")
    add_para(f"Total datasets: {ds['total']} | Total samples: {ds['samples']:,}", bold=True)
    doc.add_paragraph()

    # ── Section 6: Assessment ─────────────────────────────────────────────
    add_heading("6. Final Assessment", 1, "#3b82f6")
    add_para(a["recommendation"])
    doc.add_paragraph()
    if a["action_items"]:
        add_para("Action Items:", bold=True)
        for item in a["action_items"]:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(item).font.size = Pt(9)

    doc.add_paragraph()
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer_p.add_run(f"BlindSpot.AI v2.0  |  {report['report_id']}  |  {report['generated_at']}  |  CONFIDENTIAL")
    fr.font.size = Pt(7)
    fr.font.color.rgb = rgb("#94a3b8")

    doc.save(path)
    logger.info(f"[DOCX] Generated: {path} ({os.path.getsize(path)//1024} KB)")
